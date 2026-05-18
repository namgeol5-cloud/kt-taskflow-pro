from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import os

SCREENSHOT_DIR = os.path.join(os.path.dirname(__file__), "screenshots")
OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "taskflow-pro-test-report.docx")

TEST_CASES = [
    {
        "id": "TC01",
        "name": "초기 화면 로드 (라이트 모드)",
        "description": "앱 접속 시 라이트 테마로 초기 화면이 정상 렌더링되는지 확인",
        "steps": ["http://localhost:3000 접속", "페이지 로드 완료 대기"],
        "expected": "헤더, 업무 추가 폼, 빈 목록 메시지가 표시됨",
        "result": "PASS",
        "screenshots": ["tc01-initial-light.png"],
    },
    {
        "id": "TC02",
        "name": "업무 추가 폼 입력",
        "description": "제목·마감일·상태를 입력하고 추가 버튼 클릭 전 폼 상태 확인",
        "steps": ["제목 입력: 디자인 시안 검토", "마감일 입력: 2026-05-25 18:00", "상태 선택: In Progress"],
        "expected": "입력 필드에 값이 정상 표시됨",
        "result": "PASS",
        "screenshots": ["tc02-add-form-filled.png"],
    },
    {
        "id": "TC03",
        "name": "업무 목록 표시",
        "description": "업무 추가 후 카드 목록에 제목·D-N 카운트다운·상태 버튼이 표시되는지 확인",
        "steps": ["업무 2건 추가 완료", "목록 화면 확인"],
        "expected": "각 카드에 제목, D-N HH:MM, 상태 선택 버튼 3개가 표시됨",
        "result": "PASS",
        "screenshots": ["tc03-task-list.png"],
    },
    {
        "id": "TC04",
        "name": "상태 인라인 버튼 변경",
        "description": "카드의 Done 버튼 클릭 시 즉시 상태가 변경되는지 확인",
        "steps": ["첫 번째 카드의 Done 버튼 클릭", "API PUT 호출 및 목록 갱신 확인"],
        "expected": "Done 버튼이 초록색으로 활성화되고 상태가 즉시 반영됨",
        "result": "PASS",
        "screenshots": ["tc04-status-inline-change.png"],
    },
    {
        "id": "TC05",
        "name": "수정 모달 열기",
        "description": "카드 제목 클릭 시 수정 모달이 열리고 기존 데이터가 채워지는지 확인",
        "steps": ["카드 제목 영역 클릭", "모달 표시 확인"],
        "expected": "수정 모달이 열리고 제목·설명·마감일·상태 필드에 기존 값이 표시됨",
        "result": "PASS",
        "screenshots": ["tc05-edit-modal-open.png"],
    },
    {
        "id": "TC06",
        "name": "수정 저장",
        "description": "모달에서 제목·설명을 수정하고 저장 시 목록에 반영되는지 확인",
        "steps": [
            "제목 수정: API 명세 작성 (수정됨)",
            "설명 입력: REST API 엔드포인트 전체 문서화",
            "저장 버튼 클릭",
        ],
        "expected": "모달이 닫히고 카드에 수정된 제목·설명이 표시됨",
        "result": "PASS",
        "screenshots": ["tc06-edit-modal-filled.png", "tc06-edit-saved.png"],
    },
    {
        "id": "TC07",
        "name": "빈 제목 에러 처리 (400)",
        "description": "공백만 입력하고 추가 시 에러 메시지가 표시되는지 확인",
        "steps": ["제목에 공백만 입력", "추가 버튼 클릭"],
        "expected": "폼 하단에 '업무 제목을 입력해주세요.' 에러 메시지 표시됨",
        "result": "PASS",
        "screenshots": ["tc07-empty-title-error.png"],
    },
    {
        "id": "TC08",
        "name": "삭제 확인 다이얼로그 및 삭제",
        "description": "휴지통 아이콘 클릭 시 확인 다이얼로그가 뜨고 삭제 후 목록에서 제거되는지 확인",
        "steps": ["카드 우측 🗑 아이콘 클릭", "삭제 확인 다이얼로그 확인", "삭제 버튼 클릭"],
        "expected": "확인 다이얼로그 표시 후 삭제 시 카드가 즉시 목록에서 제거됨",
        "result": "PASS",
        "screenshots": ["tc08-delete-confirm.png", "tc08-delete-done.png"],
    },
    {
        "id": "TC09",
        "name": "다크 모드 토글",
        "description": "헤더 🌙 버튼 클릭 시 다크 테마로 전환되는지 확인",
        "steps": ["헤더 우측 테마 토글 버튼 클릭"],
        "expected": "전체 화면이 다크 테마로 전환되고 아이콘이 ☀️로 변경됨",
        "result": "PASS",
        "screenshots": ["tc09-dark-mode.png"],
    },
    {
        "id": "TC10",
        "name": "360px 모바일 반응형",
        "description": "뷰포트를 360px로 줄였을 때 레이아웃이 깨지지 않는지 확인",
        "steps": ["브라우저 뷰포트를 360×780으로 변경", "전체 페이지 확인"],
        "expected": "모든 요소가 360px 안에서 정상 배치되며 가로 스크롤 없음",
        "result": "PASS",
        "screenshots": ["tc10-mobile-360px.png"],
    },
]


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_result_badge(para, result):
    run = para.add_run(f"  [{result}]")
    run.bold = True
    run.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a) if result == "PASS" else RGBColor(0xdc, 0x26, 0x26)


def build_report():
    doc = Document()

    # 기본 스타일
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)

    # ── 표지 ──────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph("TaskFlow Pro")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(28)
    title.runs[0].bold = True

    sub = doc.add_paragraph("UI 기능 테스트 보고서")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(16)
    sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"테스트 일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    meta.add_run("테스트 환경: Windows 10 / Chrome (Playwright)\n")
    meta.add_run("테스트 대상: http://localhost:3000\n")
    meta.add_run("백엔드: FastAPI http://localhost:8000")

    doc.add_page_break()

    # ── 요약 ──────────────────────────────────────────
    add_heading(doc, "1. 테스트 요약", level=1)

    total = len(TEST_CASES)
    passed = sum(1 for tc in TEST_CASES if tc["result"] == "PASS")
    failed = total - passed

    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["총 케이스", "PASS", "FAIL", "통과율"]
    values = [str(total), str(passed), str(failed), f"{passed/total*100:.0f}%"]
    for i, (h, v) in enumerate(zip(headers, values)):
        table.cell(0, i).text = h
        table.cell(0, i).paragraphs[0].runs[0].bold = True
        c = table.cell(1, i)
        c.text = v
        run = c.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(14)
        if h == "PASS":
            run.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)
        elif h == "FAIL":
            run.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)

    doc.add_paragraph()

    # 케이스 목록 표
    add_heading(doc, "2. 테스트 케이스 목록", level=1)
    summary_table = doc.add_table(rows=1 + total, cols=4)
    summary_table.style = "Table Grid"
    for i, h in enumerate(["ID", "테스트 항목", "결과", "비고"]):
        cell = summary_table.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    for row_idx, tc in enumerate(TEST_CASES, start=1):
        row = summary_table.rows[row_idx]
        row.cells[0].text = tc["id"]
        row.cells[1].text = tc["name"]
        result_run = row.cells[2].paragraphs[0].add_run(tc["result"])
        result_run.bold = True
        result_run.font.color.rgb = (
            RGBColor(0x16, 0xa3, 0x4a) if tc["result"] == "PASS" else RGBColor(0xdc, 0x26, 0x26)
        )
        row.cells[3].text = f"{len(tc['screenshots'])}장"

    doc.add_page_break()

    # ── 각 테스트 케이스 상세 ──────────────────────────
    add_heading(doc, "3. 테스트 케이스 상세", level=1)

    for tc in TEST_CASES:
        add_heading(doc, f"{tc['id']}  {tc['name']}", level=2)

        p = doc.add_paragraph()
        add_result_badge(p, tc["result"])

        doc.add_paragraph(f"목적: {tc['description']}")

        doc.add_paragraph("테스트 절차:", style="List Bullet")
        for step in tc["steps"]:
            doc.add_paragraph(step, style="List Bullet 2")

        doc.add_paragraph(f"기대 결과: {tc['expected']}")

        for img_name in tc["screenshots"]:
            img_path = os.path.join(SCREENSHOT_DIR, img_name)
            if os.path.exists(img_path):
                caption = doc.add_paragraph(f"▶ {img_name}")
                caption.runs[0].font.size = Pt(8)
                caption.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
                doc.add_picture(img_path, width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

    doc.save(OUTPUT_PATH)
    print(f"보고서 저장 완료: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_report()
